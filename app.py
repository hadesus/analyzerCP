import os
import re
import io
import docx
import translators as ts
import requests
from bs4 import BeautifulSoup
from flask import Flask, render_template, request, redirect, url_for, flash
from werkzeug.utils import secure_filename
from flask_migrate import Migrate

from config import Config
from extensions import db
from models import Analysis, DrugResult

# --- Mappings (for normalization) ---
ROUTE_MAP = {
    'в/в': 'intravenous', 'внутривенно': 'intravenous', 'капельно': 'intravenous drip',
    'в/м': 'intramuscular', 'внутримышечно': 'intramuscular', 'п/к': 'subcutaneous',
    'подкожно': 'subcutaneous', 'перорально': 'oral', 'per os': 'oral'
}
UNITS_MAP = {
    'мг': 'mg', 'г': 'g', 'мл': 'ml', 'мкг': 'mcg', 'ме': 'iu', 'ед': 'iu', 'ui': 'iu', 'iu': 'iu'
}
# ------------------------------------

def create_app(config_class=Config):
    app = Flask(__name__)
    app.config.from_object(config_class)

    db.init_app(app)
    Migrate(app, db)

    with app.app_context():
        # Ensure instance folder and upload folder exist
        try:
            os.makedirs(app.instance_path)
            os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
        except OSError:
            pass
        # Load formulary data once on startup
        try:
            with open('who_eml_drug_list.txt', 'r', encoding='utf-8') as f:
                app.config['WHO_EML_DRUGS'] = {line.strip().lower() for line in f if line.strip()}
            print(f"--- Loaded {len(app.config['WHO_EML_DRUGS'])} drugs from formulary ---")
        except FileNotFoundError:
            app.config['WHO_EML_DRUGS'] = set()
            print("--- Formulary file not found. Please run formulary_parser.py ---")

    # --- Helper Functions ---
    def allowed_file(file):
        return file.filename and '.' in file.filename and \
               file.filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS'] and \
               file.mimetype == app.config['ALLOWED_MIMETYPE']

    # --- Core Analysis Functions ---
    def extract_drug_data_from_docx(filepath):
        try:
            document = docx.Document(filepath)
            extracted_drugs = []
            required_headers = {"inn_protocol": "международное непатентованное наименование лс", "usage_protocol": "способ применения", "loe_protocol": ["уд", "уровень доказательности"]}
            for table in document.tables:
                headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
                if required_headers["inn_protocol"] in headers and required_headers["usage_protocol"] in headers:
                    loe_header_key = next((h for h in required_headers["loe_protocol"] if h in headers), None)
                    for i in range(1, len(table.rows)):
                        cells = table.rows[i].cells
                        inn = cells[headers.index(required_headers["inn_protocol"])].text.strip()
                        if inn:
                            extracted_drugs.append({
                                "inn_protocol": inn,
                                "usage_protocol": cells[headers.index(required_headers["usage_protocol"])].text.strip(),
                                "loe_protocol": cells[headers.index(loe_header_key)].text.strip() if loe_header_key else "Н/У"
                            })
            return extracted_drugs
        except Exception as e:
            print(f"Error parsing DOCX: {e}")
            return None

    def run_full_analysis(filepath, analysis_record):
        raw_drug_list = extract_drug_data_from_docx(filepath)
        if raw_drug_list is None: return None

        disease_placeholder = "cancer"

        for raw_drug in raw_drug_list:
            drug = DrugResult(analysis_id=analysis_record.id, **raw_drug)

            usage_str = drug.usage_protocol or ""
            drug.parsed_dosage = (m.group(1).replace(',', '.') if (m := re.search(r'(\d+[\.,]?\d*)\s*(мг|г|мл|мкг|ме|ед|ui|iu)', usage_str, re.IGNORECASE)) else "Н/У")
            drug.parsed_units = (m.group(2).lower() if m else "Н/У")
            drug.parsed_frequency = (m.group(1) if (m := re.search(r'(\d+\s*раз\w*\s*в\s*\w+)', usage_str, re.IGNORECASE)) else "Н/У")
            route_str = (m.group(1).lower() if (m := re.search(r'(в/в|внутривенно|в/м|внутримышечно|п/к|подкожно|перорально|per os|капельно)', usage_str, re.IGNORECASE)) else "Н/У")
            drug.normalized_route = ROUTE_MAP.get(route_str, "unknown")

            try:
                drug.inn_english = ts.translate_text(drug.inn_protocol, to_language='en', from_language='ru')
            except Exception as e:
                print(f"Translation error for '{drug.inn_protocol}': {e}")
                drug.inn_english = "Translation Error"

            inn_eng = drug.inn_english.lower() if drug.inn_english != "Translation Error" else ""
            drug.pubmed_links = "\n".join(query_pubmed(inn_eng, disease_placeholder))
            drug.fda_status = query_regulatory_body(inn_eng, "FDA")
            drug.ema_status = query_regulatory_body(inn_eng, "EMA")
            drug.who_eml_status = "Found" if inn_eng in app.config['WHO_EML_DRUGS'] else "Not Found"
            assign_system_loe(drug)

            db.session.add(drug)

        db.session.commit()
        return analysis_record.id

    def query_pubmed(drug_name, disease):
        if not drug_name or not disease: return []
        params = {'db': 'pubmed', 'term': f'({drug_name}[Title/Abstract]) AND ({disease}[Title/Abstract]) AND (randomized controlled trial[Publication Type] OR meta-analysis[Publication Type])', 'retmode': 'json', 'retmax': 3, 'tool': app.config['PUBMED_API_TOOL'], 'email': app.config['PUBMED_API_EMAIL'], 'api_key': app.config['PUBMED_API_KEY']}
        try:
            response = requests.get("https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi", params=params)
            response.raise_for_status()
            pmids = response.json().get('esearchresult', {}).get('idlist', [])
            return [f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/" for pmid in pmids]
        except requests.RequestException as e:
            print(f"PubMed API request failed: {e}")
            return []

    def query_regulatory_body(drug_name, body):
        if not drug_name: return "N/A"
        try:
            if body == "FDA": url, selector = f"https://www.accessdata.fda.gov/scripts/cder/drugsatfda/index.cfm?fuseaction=Search.Search_Drug_Name&DrugName={drug_name}", ("table", {"id": "results-table"})
            elif body == "EMA": url, selector = f"https://www.ema.europa.eu/en/medicines/search/ema_group_types/ema_medicine?search_api_views_fulltext={drug_name}", ("div", {"class": "view-medicines-search"})
            else: return "Unknown Body"
            response = requests.get(url, headers={'User-Agent': 'Mozilla/5.0'})
            response.raise_for_status()
            return "Found" if BeautifulSoup(response.content, 'html.parser').find(*selector) else "Not Found"
        except Exception as e:
            print(f"{body} scraping failed for {drug_name}: {e}")
            return "Scraping Error"

    def assign_system_loe(drug):
        if drug.pubmed_links: drug.system_loe = "Класс I (A)"
        else: drug.system_loe = "Класс IV (D)"

    # --- Flask Routes ---
    @app.route('/')
    def index():
        return render_template('index.html')

    @app.route('/upload', methods=['POST'])
    def upload_file():
        if 'file' not in request.files:
            flash('Файл не был отправлен.', 'error')
            return redirect(url_for('index'))
        file = request.files['file']
        if not file or not file.filename:
            flash('Файл не выбран.', 'error')
            return redirect(url_for('index'))
        if allowed_file(file):
            try:
                filename = secure_filename(file.filename)
                new_analysis = Analysis(filename=filename)
                db.session.add(new_analysis)
                db.session.commit()
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], f"{new_analysis.id}_{filename}")
                file.save(filepath)
                analysis_id = run_full_analysis(filepath, new_analysis)
                if analysis_id:
                    return redirect(url_for('analysis_detail', analysis_id=analysis_id))
            except Exception as e:
                db.session.rollback()
                flash(f'Произошла непредвиденная ошибка: {e}', 'error')
        else:
            flash('Недопустимый тип или размер файла.', 'error')
        return redirect(url_for('index'))

    @app.route('/history')
    def history():
        page = request.args.get('page', 1, type=int)
        analyses = Analysis.query.order_by(Analysis.upload_timestamp.desc()).paginate(page=page, per_page=10)
        return render_template('history.html', analyses=analyses)

    @app.route('/analysis/<int:analysis_id>')
    def analysis_detail(analysis_id):
        analysis = Analysis.query.get_or_404(analysis_id)
        return render_template('analysis_detail.html', analysis=analysis)

    @app.route('/export/<int:analysis_id>')
    def export_results(analysis_id):
        analysis = Analysis.query.get_or_404(analysis_id)
        document = docx.Document()
        document.add_heading(f'Результаты Анализа: {analysis.filename}', 0)
        headers = ['Название (из протокола)', 'МНН (ENG)', 'Применение (из протокола)', 'Статус в источниках', 'Ссылки PubMed', 'УД (из протокола)', 'Системный УД']
        table = document.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
        for drug in analysis.drug_results:
            row_cells = table.add_row().cells
            row_cells[0].text = drug.inn_protocol or ''
            row_cells[1].text = drug.inn_english or ''
            row_cells[2].text = drug.usage_protocol or ''
            row_cells[3].text = f"WHO EML: {drug.who_eml_status}\nFDA: {drug.fda_status}\nEMA: {drug.ema_status}"
            row_cells[4].text = drug.pubmed_links if drug.pubmed_links else "Нет"
            row_cells[5].text = drug.loe_protocol or ''
            row_cells[6].text = drug.system_loe or ''
        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        return send_file(file_stream, as_attachment=True, download_name=f'report_{analysis.id}.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    return app

app = create_app()
if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=8080)
